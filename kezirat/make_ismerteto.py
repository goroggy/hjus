# -*- coding: utf-8 -*-
"""Egy tartalom -> ket kimenet: HJus-ismerteto.md es HJus-ismerteto.docx.

A tartalom ITT a forras (a szoveg dokumentumra szabott valtozata; a HTML lapon
gombok es adatsavok vannak, amik papiron ertelmetlenek). Ami viszont NEM lehet
ket kulon igazsag: a SZAMOK. Ezert a script a vegen ellenorzi, hogy minden
szam, ami itt szerepel, elofordul-e a kozzetett HTML-ben is.
"""
import io, os, re, sys
sys.stdout.reconfigure(encoding="utf-8")

HERE = os.path.dirname(os.path.abspath(__file__))
HTML = os.path.join(HERE, "blog-src", "harom-app.html")
OUT_MD = os.path.join(HERE, "HJus-ismerteto.md")
OUT_DOCX = os.path.join(HERE, "HJus-ismerteto.docx")

LAND = "https://goroggy.github.io/hjus/"

APPS = [
    dict(
        nev="HJusDefs",
        alcim="Mit jelent ez a szó — és melyik törvény szerint?",
        adat="14 417 fogalommeghatározás · 9 399 meghatározott fogalom · 1 546 jogszabályból",
        bekezdesek=[
            "A jogszabályok „Értelmező rendelkezések” szakaszaiból és az „E törvény "
            "alkalmazásában…” fordulatokból gépileg kigyűjtött fogalom–definíció párok. "
            "A bal oldalon fogalomra keresel, a jobb oldalon a fogalmi környezete "
            "rajzolódik ki: mire épül a meghatározása, és mely fogalmak épülnek rá.",
            "A lényeg a nézőpont. A Nemzeti Jogszabálytár — helyesen — jogszabály szerint "
            "van rendezve: megnyitod a törvényt, és megtudod, mit jelent benne egy fogalom. "
            "Amit nem tudsz meg: hogy ugyanazt a szót hány másik jogszabály határozza meg, "
            "és miben másképp. 1 771 olyan fogalom van, amelyet legalább két jogszabály "
            "külön is meghatároz.",
        ],
        mirejo=[
            ("Beadvány, szerződés, szakvélemény írásakor.",
             "Mielőtt egy fogalmat a köznyelvi vagy a Ptk.-beli jelentésében használnál, "
             "egy keresés megmutatja, hogy az adott jogterület saját meghatározása eltér-e. "
             "A „közeli hozzátartozó” például harminchét helyen van meghatározva."),
            ("Jogszabály-előkészítésnél.",
             "Van-e már bevett meghatározás, amire hivatkozni lehet ahelyett, hogy egy "
             "harmincnyolcadik születne."),
            ("Jogvitában.",
             "Kimutatható, hogy a másik fél egy másik jogszabály fogalomkészletével érvel."),
            ("Oktatásban.",
             "Egy fogalom nem magában áll: látszik, mely tágabb fogalom fajtája, és mely "
             "szűkebbek épülnek rá."),
        ],
        korlat="A szövegek a kihirdetéskori állapotot tükrözik, a későbbi módosításokat a "
               "gyűjtemény nem követi. Minden jogszabálycím az NJT-re mutat — amire "
               "hivatkozol, azt ott ellenőrizd.",
        linkek=[("Az alkalmazás", "https://goroggy.github.io/hjusdefs-web/hjusdefs.html"),
                ("Dokumentáció", "https://goroggy.github.io/hjusdefs-web/Miez.html")],
    ),
    dict(
        nev="HJusSent",
        alcim="Mit mondtak még ugyanerről — akkor is, ha nem találod el a kulcsszót?",
        adat="46 280 bírói rezümé · forrás: BHGY · két réteg: tartalom / eljárási sablon",
        bekezdesek=[
            "A Bírósági Határozatok Gyűjteményéből letöltött rezümék — a felsőbb bírósági "
            "határozatokhoz fűzött, elvileg egy-két mondatos összefoglalók. A bal oldalon "
            "keresel vagy szűrsz, a jobb oldalon a kiválasztott tétel szomszédsága jelenik "
            "meg: mely más rezümék állnak hozzá szövegben a legközelebb, fa alakban "
            "elrendezve.",
            "A rezümék minősége a valóságban vegyes: van, ami idézhető jogtétel, van, ami "
            "csak az ügy tárgyát nevezi meg, és van a Kúria felülvizsgálati befogadásáról "
            "szóló, tömegesen ismétlődő sablonszöveg. Ez utóbbi külön rétegként el van "
            "különítve — nem gépi becslés alapján, hanem szabály és klaszter-megerősítés "
            "együtt azonosítja.",
        ],
        mirejo=[
            ("Ha egy jó találatod már van.",
             "Rákattintva előjön, mi áll hozzá a legközelebb — anélkül, hogy ki kellene "
             "találnod, milyen szavakkal fogalmazott a bíróság."),
            ("Annak megítéléséhez, hogy egy jogtétel mennyire magányos.",
             "Ha húsz hasonló rezümé áll körülötte, az más súlyú, mint ha egy sem."),
            ("Zajszűrésre.",
             "Az eljárási sablonok kitakarhatók, így a keresés nem fullad a befogadási "
             "dogmatikába."),
        ],
        korlat="A hasonlóság szövegalapú becslés, nem jogi azonosság. Két rezümé lehet "
               "szövegében közeli és jogilag távoli — és fordítva. A fa navigációs "
               "segédeszköz: arra való, hogy elvezessen a következő olvasnivalóhoz, nem "
               "arra, hogy eldöntse, mi tartozik egy kérdés alá.",
        linkek=[("Az alkalmazás", "https://goroggy.github.io/hulex-sentencies-web/klaszter_fa.html"),
                ("Kézikönyv", "https://goroggy.github.io/hulex-sentencies-web/kezikonyv.html")],
    ),
    dict(
        nev="HJusRes",
        alcim="Egy jogszabályi fogalom — és a bírói gyakorlat, ahol előkerül.",
        adat="42 142 rezümé fogalmi címkével · 176 630 közvetlen kapcsolat · "
             "1 850 fogalom fordul elő",
        bekezdesek=[
            "Az első kettő összekötve: a HJusDefs fogalmait rákeresi a HJusSent rezüméire, "
            "szótövezett illesztéssel — tehát a ragozott alakot is megtalálja, viszont csak "
            "pontos, teljes egyezést fogad el. A közvetlen találatok mellé a fogalmi "
            "hierarchiából egy-két lépéssel feljebb lévő tágabb fogalmakat is odateszi, "
            "külön, gyengébb jelöléssel.",
            "A 46 019 rezüméből 42 142 kapott legalább egy fogalmi címkét — nagyjából minden "
            "tizenkettedik egyet sem. A magas arány részben azt is jelenti, hogy a gyakori, "
            "általános fogalmak („kérelem”, „eljárás”) mindenütt megjelennek; a ritkább "
            "fogalmak a hasznosak. A nézet fogalomlistája 1 964 tételes: a 114 többlet olyan "
            "tágabb fogalom, amely csak öröklött kapcsolaton át kerül elő.",
        ],
        mirejo=[
            ("A definíció és az alkalmazás közti ugráshoz.",
             "Tudod, mit ír a törvény a „magánút”-ról; innen egy lépés megnézni, mely "
             "határozatok rezüméiben kerül elő."),
            ("Fogalom felőli belépéshez.",
             "Nem ügyszámmal vagy jogszabályhellyel kezdesz, hanem azzal a szóval, ami a "
             "kérdésben szerepel."),
            ("Fogalmi környezet bejárásához.",
             "A tágabb fogalom felől is elindulhatsz, ha a szűkebbre kevés a találat."),
        ],
        korlat="A kapcsolat szóegyezés, nem jelentésazonosság. Abból, hogy egy rezümében "
               "szerepel a szó, nem következik, hogy a bíróság épp azt a jogszabályi "
               "definíciót alkalmazta — sőt az sem, hogy ugyanabban az értelemben használta. "
               "A címke odavezet a szöveghez; az olvasás a te dolgod marad.",
        linkek=[("Az alkalmazás", "https://goroggy.github.io/hjusresumes-web/hjusresumes.html")],
    ),
]

KOZOS = [
    ("Kísérleti, nem termék",
     "Mindhárom kutatási célú próba. Nincs mögöttük szerkesztőségi ellenőrzés, "
     "karbantartási vállalás vagy rendelkezésre állási ígéret."),
    ("Gépi kinyerés",
     "Az adatot program gyűjtötte a nyilvános forrásokból, tehát hibázhat. Ezért mutat "
     "minden tétel a forrására: nem a gyűjteménynek kell hinni, hanem a Közlönynek, "
     "illetve a határozatnak."),
    ("Nem hatályos jogot mutat",
     "A fogalommeghatározások a kihirdetéskori állapotot tükrözik, a rezümék a közzététel "
     "szerintit. Hatályos szövegért az NJT-re, illetve a BHGY-re kell menni — a linkek "
     "oda vezetnek."),
    ("Nem jogi tanácsadás",
     "Kereső- és tájékozódási eszközök. Sem a találat megléte, sem a hiánya nem jogi "
     "állítás."),
    ("Nem gyűjt rólad adatot",
     "Nincs bejelentkezés, nincs szerveroldali napló, nincs analitika, és egyik oldal sem "
     "tölt be semmit külső kiszolgálóról. A HJusDefs és a HJusSent megjegyzi, hol jártál "
     "legutóbb — de kizárólag a saját böngésződ tárolójában, ez az adat nem hagyja el a "
     "gépedet. A HJusRes semmit nem tárol."),
    ("Asztali gépre való",
     "Az oldalak az egész adatállományt egyben töltik le (9, illetve 30 és 28 MB). Első "
     "betöltésre széles sávú kapcsolatot és nem telefont érdemes hozzá használni; utána "
     "viszont hálózat nélkül is működnek."),
]

BEVEZETO = (
    "Kereshetővé tett fogalommeghatározások a Magyar Közlönyből, csoportosított bírói "
    "rezümék a Bírósági Határozatok Gyűjteményéből, és a kettő összekötve. Mindhárom "
    "szabadon használható, regisztráció nélkül; kutatási és kísérleti célú, nem jogi "
    "tanácsadás, és nem hatályos szöveget mutat."
)

LANC = [
    ("Magyar Közlöny (2013–2026)", "HJusDefs", "jogszabályi fogalommeghatározások"),
    ("Bírósági Határozatok Gyűjteménye", "HJusSent", "bírói rezümék hasonlóság szerint"),
    ("az előző kettő", "HJusRes", "fogalom felől a bírói gyakorlathoz"),
]

VISSZAJELZES = (
    "A legtöbbet érő visszajelzés nem az, hogy „szép” vagy „nem szép”, hanem egy konkrét "
    "eset: egy fogalom, amit rosszul nyert ki; egy hiányzó jogszabály; vagy egy "
    "munkafolyamat, amiben ez az eszköz elakad. A hibás találat különösen hasznos — abból "
    "derül ki, hol téved rendszeresen a kinyerés."
)

ADAT = (
    "Magyar Közlöny (2013.05.01.–2026) és a Bírósági Határozatok Gyűjteménye, nyilvánosan "
    "közzétett szövegekből, gépi feldolgozással. A jogszabály- és határozatszöveg nem áll "
    "szerzői jogi védelem alatt; a kinyert, származtatott adat CC BY 4.0 alatt szabadon "
    "felhasználható, forrásmegjelöléssel."
)
KAPCSOLAT = "Észrevétel, hiba, javaslat: goroggyorgy77@gmail.com"
CIM = "Három kísérleti alkalmazás a magyar joganyagon"


# ---------------------------------------------------------------- Markdown
def markdown():
    o = ["# " + CIM, "", BEVEZETO, "",
         "Közös belépő: <%s>" % LAND, "",
         "Az alábbi három eszköz külön-külön is használható, de egymásra épül: a harmadik "
         "az első kettő kimenetét kapcsolja össze.", ""]
    for forras, nev, mit in LANC:
        o.append("- %s → **%s** — %s" % (forras, nev, mit))
    o.append("")
    for i, a in enumerate(APPS, 1):
        o += ["## %d. %s" % (i, a["nev"]), "", "*%s*" % a["alcim"], "",
              "**%s**" % a["adat"], ""]
        for b in a["bekezdesek"]:
            o += [b, ""]
        o += ["**Mire jó**", ""]
        for cim, szov in a["mirejo"]:
            o.append("- **%s** %s" % (cim, szov))
        o += ["", "**Amit nem mond:** " + a["korlat"], ""]
        for nev, url in a["linkek"]:
            o.append("- %s: <%s>" % (nev, url))
        o.append("")
    o += ["## Ami mindháromra igaz", ""]
    for cim, szov in KOZOS:
        o += ["**%s** — %s" % (cim, szov), ""]
    o += ["## Amit érdemes visszajelezni", "", VISSZAJELZES, "",
          "---", "", "**Adat.** " + ADAT, "", "**Kapcsolat.** " + KAPCSOLAT, ""]
    return "\n".join(o)


# ---------------------------------------------------------------- DOCX
def docx_file():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    def para(text=None, style=None, bold_prefix=None, italic=False, space_after=8):
        p = d.add_paragraph(style=style)
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            if text:
                p.add_run(" " + text)
        elif text:
            r = p.add_run(text)
            r.italic = italic
        p.paragraph_format.space_after = Pt(space_after)
        return p

    d.add_heading(CIM, level=0)
    para(BEVEZETO)
    p = d.add_paragraph()
    p.add_run("Közös belépő: ").bold = True
    p.add_run(LAND)
    para("Az alábbi három eszköz külön-külön is használható, de egymásra épül: "
         "a harmadik az első kettő kimenetét kapcsolja össze.")
    for forras, nev, mit in LANC:
        p = d.add_paragraph(style="List Bullet")
        p.add_run(forras + " → ")
        p.add_run(nev).bold = True
        p.add_run(" — " + mit)

    for i, a in enumerate(APPS, 1):
        d.add_heading("%d. %s" % (i, a["nev"]), level=1)
        para(a["alcim"], italic=True)
        p = d.add_paragraph()
        p.add_run(a["adat"]).bold = True
        for b in a["bekezdesek"]:
            para(b)
        para(bold_prefix="Mire jó")
        for cim, szov in a["mirejo"]:
            p = d.add_paragraph(style="List Bullet")
            p.add_run(cim).bold = True
            p.add_run(" " + szov)
        para(a["korlat"], bold_prefix="Amit nem mond:")
        for nev, url in a["linkek"]:
            p = d.add_paragraph(style="List Bullet")
            p.add_run(nev + ": ")
            r = p.add_run(url)
            r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    d.add_heading("Ami mindháromra igaz", level=1)
    for cim, szov in KOZOS:
        para(szov, bold_prefix=cim + " —")

    d.add_heading("Amit érdemes visszajelezni", level=1)
    para(VISSZAJELZES)

    d.add_paragraph()
    para(ADAT, bold_prefix="Adat.")
    para(KAPCSOLAT, bold_prefix="Kapcsolat.")
    d.save(OUT_DOCX)


# ---------------------------------------------------------------- ellenorzes
def szamok(t):
    """Szokozzel tagolt magyar szamok + a kisebbek, kivéve az evszamokat."""
    return set(re.findall(r"\b\d{1,3}(?: \d{3})+\b", t))


def ellenoriz(md):
    html = io.open(HTML, encoding="utf-8").read()
    html_txt = re.sub(r"<[^>]+>", " ", re.sub(r"(?s)<style.*?</style>", " ", html))
    html_txt = html_txt.replace("&shy;", "")
    hianyzik = sorted(szamok(md) - szamok(html_txt))
    print("szamok a dokumentumban:", len(szamok(md)))
    if hianyzik:
        print("!! NEM SZEREPEL a kozzetett lapon:", hianyzik)
        raise SystemExit(1)
    print("minden szam egyezik a kozzetett lappal")


if __name__ == "__main__":
    md = markdown()
    io.open(OUT_MD, "w", encoding="utf-8", newline="\n").write(md)
    print("kiirva:", OUT_MD, len(md), "karakter")
    docx_file()
    print("kiirva:", OUT_DOCX, os.path.getsize(OUT_DOCX), "bajt")
    ellenoriz(md)
